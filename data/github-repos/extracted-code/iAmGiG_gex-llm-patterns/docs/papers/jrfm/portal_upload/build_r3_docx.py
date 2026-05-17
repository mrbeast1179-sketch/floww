"""Generate the Reviewer 3 response as a .docx that follows the MDPI
response-to-reviewer template structure.

The MDPI template (shipped alongside this file as
``Example for author to respond reviewer - MDPI.docx``) requires five
numbered sections:

    1. Summary (neutral thank-you)
    2. Questions for General Evaluation (table of ratings + responses)
    3. Point-by-point response to Comments and Suggestions for Authors
       (Comments N: / Response N: format)
    4. Response to Comments on the Quality of English Language
    5. Additional clarifications

Our earlier point-by-point response (response_to_reviewers.md) used
internal R3.1-R3.9 numbering and did not include sections 1, 2, 4, or 5.
This script builds a .docx with the correct structure, with the
reviewer's verbatim comments pasted into each Comments N: entry and our
response text pasted into each Response N: entry.

Usage:
    python build_r3_docx.py

Outputs:
    response_R3_MDPI.docx  (upload this to the portal)

Requires: python-docx (installed via pip).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
import os

# Default output; if the file is locked (likely open in Word), save to a
# sibling with a _v2 suffix so the user can diff and replace manually.
PRIMARY = HERE / "response_R3_MDPI.docx"
FALLBACK = HERE / "response_R3_MDPI_v2.docx"


def _select_output() -> Path:
    try:
        # Test write-ability: open the primary for append-binary and close.
        if PRIMARY.exists():
            with PRIMARY.open("ab"):
                pass
        return PRIMARY
    except (OSError, PermissionError):
        return FALLBACK


OUTPUT = _select_output()


# ---------- Content: reviewer comments (verbatim, in order of appearance) ----------

REVIEWER_SUMMARY_COMMENT = (
    "The manuscript addresses a timely and interesting topic at the "
    "intersection of financial market microstructure and large language "
    "model validation. The idea of using temporal obfuscation to "
    "distinguish structural reasoning from memorization is original and "
    "potentially valuable. However, several aspects of the paper would "
    "benefit from further clarification, strengthening, and refinement."
)

# Per-item reviewer comments (8 substantive items). Paragraphs are
# transcribed verbatim from the review report as received.
COMMENTS = [
    (
        "1",
        "The introduction must be shortened and made more focused. It currently "
        "contains overly long and philosophical paragraphs. It should clearly "
        "state the research gap, the contribution, and how the paper differs "
        "from existing studies in financial econometrics. More recent references "
        "(especially 2022-2025) on options market microstructure, gamma "
        "exposure, and 0DTE dynamics must be added and critically discussed.",
    ),
    (
        "2",
        "The positioning of the paper must be clarified. It is not clear "
        "whether the contribution is mainly methodological (LLM validation) or "
        "financial (market microstructure). This needs to be explicitly stated "
        "and consistently reflected throughout the paper.",
    ),
    (
        "3",
        "The research design must be strengthened. The paper currently lacks "
        "comparison with standard benchmark models such as regime-switching "
        "models or volatility-based approaches. At least one benchmark model "
        "should be included to validate the added value of the proposed "
        "framework. The causal interpretation related to 0DTE should be "
        "moderated or supported with stronger empirical evidence.",
    ),
    (
        "4",
        "The methodology section needs more transparency. The exact prompts "
        "used for the LLM must be provided (preferably in an appendix). The "
        "choice of thresholds (70% persistence, $5B magnitude, <=5 flips) must "
        "be justified or tested through sensitivity analysis. The impact of "
        "model parameters (e.g., temperature = 1.0) on reproducibility must be "
        "explained.",
    ),
    (
        "5",
        "The results section must include statistical validation. The paper "
        "relies heavily on percentages without reporting statistical "
        "significance, confidence intervals, or robustness tests. These must "
        "be added. Some interpretations are too strong compared to the evidence "
        "and should be moderated.",
    ),
    (
        "6",
        "The discussion must be better connected to finance. The implications "
        "for risk management, market efficiency, and practitioners should be "
        "explicitly developed. The current discussion is too general and "
        "sometimes theoretical.",
    ),
    (
        "7",
        "The limitations section must be expanded. It should clearly address "
        "the use of a single asset (SPY), the dependence on one LLM model, and "
        "the lack of external validation.",
    ),
    (
        "8",
        "Figures and tables must be improved. Some are too dense and difficult "
        "to read. Labels and captions should be clearer and more explanatory.",
    ),
]

ENGLISH_COMMENT = (
    "The clarity of the manuscript needs improvement. Many sentences are too "
    "long and complex, which affects readability. The writing should be "
    "simplified by using shorter sentences, more direct wording, and by "
    "removing redundant or overly elaborate expressions. Careful language "
    "editing is recommended to improve clarity and flow."
)


# ---------- Content: our responses (paragraph-mode text) ----------

# Each entry is a list of paragraphs; each paragraph is a list of
# (text, bold) tuples. bold=True renders the run in bold; bold=False
# renders plain. Response text is rendered in RED per the MDPI template
# convention.

R: dict[str, list[list[tuple[str, bool]]]] = {
    "1": [
        [
            ("We agree and have rewritten Section 1 Introduction to address each element of this comment.", False),
        ],
        [
            ("(i) Shortened and less philosophical. ", True),
            (
                'The original paragraph-1 opener ("The decisive question '
                'confronting any deployment of large language models...") '
                "has been removed. The new Section 1 opens with a two-sentence, "
                "direct statement of the validation problem and why it is "
                "first-order in finance specifically.",
                False,
            ),
        ],
        [
            ("(ii) Explicit research gap. ", True),
            (
                'A new paragraph titled "Research gap" follows the opener. '
                "It names what prior literature has done independently "
                "(dealer-gamma microstructure, 0DTE growth, LLM-reasoning "
                "probing in non-financial domains) and states precisely which "
                "combination has not been attempted: an LLM structural-reasoning "
                "validation method that (a) controls for training-data "
                "memorisation of specific events and dates, (b) is tested at a "
                "scale comparable to the target domain, and (c) discriminates "
                "genuine structural detection from reproduction of a "
                "volatility-regime classifier. The Markov-switching benchmark "
                "added per the reviewer's comment 3 below is then introduced "
                "as the direct test of element (c).",
                False,
            ),
        ],
        [
            ("(iii) Why 0DTE matters here. ", True),
            (
                'A new "Why 0DTE matters here" paragraph replaces the previous '
                '"practical urgency" framing. It explains that 0DTE growth is '
                "a natural setting for an obfuscation study because it created "
                "an observable structural shift within the training horizon of "
                "modern LLMs.",
                False,
            ),
        ],
        [
            ("(iv) 2022-2025 references added and critically discussed. ", True),
            (
                "The key new citation is Dim, Eraker & Vilkov (2023) "
                '("0DTEs: Trading, Gamma Risk and Volatility Propagation", '
                "SSRN 4692190), which provides the first systematic empirical "
                "study of 0DTE dealer inventory. It is now cited in Section 1 "
                "and critically discussed in Section 2.2, noting that it "
                "establishes dealer-hedging rather than information flow as "
                "the dominant channel through which 0DTE trading affects the "
                "underlying. We retain the existing 2022-2025 refs "
                "(Anderegg et al. 2022; Fishman 2023 Goldman Sachs; "
                "CBOE 2024 and 2025 research notes; Dim, Marsh, Schrimpf 2025 "
                "BIS).",
                False,
            ),
        ],
        [
            (
                "Change location: Section 1 Introduction paragraphs 1-4 (full rewrite), Section 2.2 Zero-Days-to-Expiration Options (new critical discussion of dim2023odtes), references.bib (new dim2023odtes entry).",
                True,
            ),
        ],
    ],
    "2": [
        [
            (
                "We agree and have stated the positioning explicitly in two "
                "places to ensure the stance is consistent throughout the "
                "paper.",
                False,
            ),
        ],
        [
            (
                "The primary contribution is methodological: temporal "
                "obfuscation testing (with the WHO->WHOM->WHAT causal "
                "framework and multi-scale validation protocol) as a "
                "generalizable procedure for validating LLM structural "
                "reasoning. Options dealer gamma-exposure regime detection is "
                "the empirical demonstration domain, selected because it "
                "combines theoretically grounded mechanical constraints, a "
                "large quantitative testbed, and the sharp pre- vs "
                "post-0DTE temporal contrast. The financial-market findings "
                "(69.1pp detection gap, 0% false-positive rate on synthetic "
                "controls, 2021-2024 0DTE-tracking regime evolution) are "
                "downstream evidence that the methodology discriminates "
                "correctly, not novel claims about options microstructure.",
                False,
            ),
        ],
        [
            (
                'Change location: new Section 1.3 "Positioning" subsection (between Contributions and Paper Organization) and Section 7 Conclusion opening (rewritten to echo the same stance before the four numbered contributions).',
                True,
            ),
        ],
    ],
    "3": [
        [
            (
                "We have added a two-state Markov-switching regression benchmark and moderated the 0DTE causal language in parallel.",
                False,
            ),
        ],
        [
            ("(a) Benchmark comparison. ", True),
            (
                "We fit statsmodels.tsa.regime_switching.MarkovRegression "
                "(2-state, switching intercept and variance, standard EM) to "
                "(i) SPY daily log returns for 2020 (canonical volatility "
                "benchmark), (ii) SPY daily log returns for 2024 (same), and "
                "(iii) the 2024 daily net-GEX series (GEX-native analogue "
                "benchmark). Per-window agreement with LLM labels: 2020 "
                "returns N=201, kappa=0.045; 2024 returns N=222, kappa=-0.178; "
                "2024 net-GEX N=221, kappa=0.610. The LLM detector is not "
                "reducible to a returns-based volatility regime (kappa near 0 "
                "or negative) but is consistent with a mechanical 2-state "
                "Gaussian on the same physical series (substantial "
                "agreement). This directly answers the reviewer's concern: "
                "the LLM reasons about dealer-gamma structure, not variance "
                "regimes.",
                False,
            ),
        ],
        [
            ("(b) Moderated 0DTE causal language. ", True),
            (
                'Section 6.3 "Market Structure Evolution and 0DTE Hypothesis" '
                "has been rewritten with explicit causal-inference hygiene: "
                "the 0DTE correspondence is framed as temporal coincidence "
                "supported by a plausible mechanical channel rather than a "
                "demonstrated causal relationship; four concurrent "
                "confounders are named (interest rates, systematic short-vol "
                "flow, passive/index AUM growth, market-maker concentration); "
                "three candidate causal-identification designs are proposed "
                "(a 0DTE suspension natural experiment, a counterfactual "
                "non-SPY launch, an instrumental-variable design); and the "
                'discussion closes with the explicit caveat that "less '
                'easily reconciled" is not "ruled out". Section 6 '
                "Conclusion contribution 3 similarly replaces "
                '"0DTE-driven structural reorganization" with '
                "temporal-coincidence language.",
                False,
            ),
        ],
        [
            (
                "Change location: new Section 3.8 Markov-Switching Benchmark, new Section 5.6 Comparison with Markov-Switching Benchmark (Table 6 + Figure 8), Section 6.3 rewrite, Section 7 Conclusion contribution 3 rewrite.",
                True,
            ),
        ],
    ],
    "4": [
        [
            ("We have addressed this comment in three parts.", False),
        ],
        [
            ("(a) Exact prompts. ", True),
            (
                "The complete regime-detection prompt is now reproduced "
                "verbatim in a new Appendix A, together with the actual "
                "OpenAI Batch API configuration we used (model o4-mini; "
                "temperature defaults to 1 because reasoning models reject "
                "user-supplied temperature overrides; max_completion_tokens "
                "not explicitly set, so the OpenAI API default applies; JSON "
                "structure requested in the prompt rather than enforced via "
                "the response_format field) and the output JSON schema used "
                "for parsing. The appendix is transcribed directly from the "
                "build_regime_prompt() function in the publicly released "
                "source code.",
                False,
            ),
        ],
        [
            ("(b) Threshold sensitivity. ", True),
            (
                "A 5×3×3 grid sweep (persistence in {60, 65, 70, "
                "75, 80}%, magnitude in {$3B, $5B, $7B}, flips <= {3, 5, 7}; "
                "45 configurations) has been applied to the 223 Phase 3 "
                "(2024) and 220 Phase 4 (2020) per-window records already "
                "on disk. Results: the 2024-vs-2020 detection gap ranges "
                "[34.1, 85.2] pp across configurations (median 63.2 pp) and "
                "exceeds 50 pp in 40/45 configurations. Reported in new "
                "Section 5.5 Threshold Sensitivity with Figure 7 heatmap.",
                False,
            ),
        ],
        [
            ("(c) Temperature / reproducibility. ", True),
            (
                "Appendix A contains a Reproducibility note explaining that "
                "OpenAI reasoning models (o1, o3, o4-mini, and GPT-5 "
                "reasoning variants) reject user-supplied temperature / "
                "top_p values and run at the default temperature of 1. The "
                "seed parameter is supported by o4-mini (OpenAI documents "
                "it as best-effort determinism that can shift when the "
                "server system_fingerprint changes), but we did not set a "
                "seed in this study. Bit-identical reproduction of any "
                "single response is therefore not guaranteed. "
                "Reproducibility at the distributional level is established "
                "through the N = 2,221 evaluation sample and the mechanical "
                "numerical thresholds embedded in the prompt itself.",
                False,
            ),
        ],
        [
            (
                "Change location: new Appendix A (pp. 24-29 in the revised PDF), new Section 5.5 Threshold Sensitivity, cross-reference added in Section 3.5 LLM Configuration pointing to Appendix A.",
                True,
            ),
        ],
    ],
    "5": [
        [
            ("We agree. The revision addresses this comment in four parts.", False),
        ],
        [
            ("(a) Confidence intervals. ", True),
            (
                "Every detection rate reported in Section 4 Results now "
                "carries a 95% confidence interval. For Phases 1-4 and all "
                "Phase 2 negative controls we report a 10,000-replicate "
                "percentile bootstrap over windows (deterministic seed); for "
                "Phase 5 per-year rates we report 95% Wilson score intervals "
                "(Brown, Cai & DasGupta, 2001). Phase 3 full 2024: 81.2% "
                "[75.8, 86.1]%. Phase 4 full 2020: 12.1% [8.1, 16.6]%. The "
                "2020 upper CI bound (17.3%) does not overlap the 2024 lower "
                "CI bound (75.8%), which directly supports the 69.1 pp "
                "separation claim with bounded evidence rather than point "
                'estimates alone. A new "Statistical conventions" paragraph '
                "at the head of Section 4.1 documents the methodology.",
                False,
            ),
        ],
        [
            ("(b) Expanded chi-square / Fisher reporting. ", True),
            (
                "Phase 4 (2020 vs 2024): Pearson's chi-square = 213.67 "
                "(df=1, p = 2.2e-48), Yates-corrected chi-square = 210.90 "
                "(p = 8.7e-48), Fisher's exact two-sided p = 1.8e-52 "
                "(odds ratio 31.3), phi = 0.69, risk difference 69.1 pp "
                "(95% Wald CI [62.4, 75.7] pp). Phase 5 (2023 -> 2024 "
                "transition): chi-square = 314.4 (p = 2.4e-70), Fisher's "
                "exact p = 9.9e-87, phi = 0.82. Abstract and Introduction "
                "updated to report Fisher's exact p rather than a bare "
                '"p < 0.0001".',
                False,
            ),
        ],
        [
            ("(c) Robustness. ", True),
            (
                "The 45-configuration threshold-sensitivity sweep described "
                "under comment 4(b) above functions as the robustness test "
                "(gap > 50 pp in 40/45 configurations).",
                False,
            ),
        ],
        [
            ("(d) Moderated claim language. ", True),
            (
                "Section 7 Conclusion contribution 2 now reports the 69.1 pp "
                "separation with explicit CI brackets on each rate and "
                "Fisher's exact p, and cites the 45-configuration robustness "
                "of the 50 pp gap. Contribution 3 moderates the 0DTE-causal "
                "language (see comment 3(b) above). Section 6.3 similarly "
                'softens "tipping-point dynamic strengthens the structural '
                'interpretation" to "is consistent with, rather than proof '
                'of". Statistical claims on the 2020-vs-2024 separation are '
                "preserved as-is; only the causal-inference language around "
                "0DTE is moderated.",
                False,
            ),
        ],
        [
            (
                "Change location: Section 4.1 statistical conventions paragraph; Section 4.3 Phase 1/3 inline CIs; Tables 2, 3, 4, 5 CI columns; references.bib added brown2001interval; new reprocessing scripts under scripts/validation/paper2/jrfm_revision/.",
                True,
            ),
        ],
    ],
    "6": [
        [
            ("We agree that the original discussion was too general on the practitioner side.", False),
        ],
        [
            (
                'The previous Section 6.6 "Practitioner Implications" has '
                'been renamed "Practical Implications" and restructured '
                "into three explicit subsubsections matching the three axes "
                "the reviewer identified:",
                False,
            ),
        ],
        [
            ("Risk management. ", True),
            (
                "Three concrete applications developed: intraday volatility "
                "budgeting (regime as leading indicator for volatility-of-"
                "volatility exposure sizing), option-book hedging under OpEx "
                "concentration, and risk-scenario design (2020 fragmented vs "
                "2024 persistent-negative as natural conditioning variables).",
                False,
            ),
        ],
        [
            ("Market efficiency. ", True),
            (
                "A positive account is offered: the detection-alpha "
                "orthogonality is consistent with a weakly efficient market "
                "in which structural constraints are reliably identifiable "
                "but already priced. This reconciles persistent microstructure "
                "influence with Sharpe deterioration.",
                False,
            ),
        ],
        [
            ("Practitioners: pipeline design and deployment. ", True),
            (
                "Two design implications developed: (i) the 30.8 pp advantage "
                "of raw strike-level data over pre-aggregated GEX challenges "
                "the default of parametric aggregation, with generalisations "
                "to credit risk, fixed-income surveillance, and equity factor "
                "research; (ii) the 2022-2024 0DTE regime shift implies that "
                "static microstructure models calibrated to pre-2022 data "
                "need recalibration.",
                False,
            ),
        ],
        [
            (
                'Change location: Section 6.6 "Practical Implications" (renamed from "Practitioner Implications"), three new subsubsections.',
                True,
            ),
        ],
    ],
    "7": [
        [
            ("We thank the reviewer for flagging these specific omissions.", False),
        ],
        [
            (
                'Section 6.7 has been renamed "Limitations and Future Work" '
                "and expanded from six limitations to seven. Each item is "
                "now explicitly tied to a concrete follow-up study. The "
                "three items the reviewer named are now addressed as:",
                False,
            ),
        ],
        [
            ("(a) Single-asset scope. ", True),
            (
                'Item 1 ("Single-asset scope") explicitly acknowledges '
                "that all results concern SPY, lists QQQ, IWM, individual "
                "equities, and non-equity underliers as relevant but "
                "untested targets, and identifies cross-asset replication as "
                "the single highest-priority item for future work.",
                False,
            ),
        ],
        [
            ("(b) Single-LLM dependence. ", True),
            (
                "A dedicated second item proposes a model-swap protocol "
                "covering Anthropic Claude, OpenAI o3, Google Gemini, and "
                "open-source reasoning models using identical prompts, with "
                "cross-model agreement analysis as the diagnostic.",
                False,
            ),
        ],
        [
            ("(c) Lack of independent external validation. ", True),
            (
                "A new third item acknowledges that per-window ground-truth "
                "metrics are computed from the same Alpha Vantage feed used "
                "to construct the windows, and proposes cross-validation "
                "against CBOE DataShop / OPRA / commercial vendors "
                "(SpotGamma, MenthorQ) and against related microstructure "
                "observables.",
                False,
            ),
        ],
        [
            (
                "Change location: Section 6.7 Limitations and Future Work (renamed, expanded 6 -> 7 items, each with explicit future-work sentence).",
                True,
            ),
        ],
    ],
    "8": [
        [
            (
                "We addressed the comment on figures and tables in two "
                "complementary ways: (i) a figure-font pass to raise "
                "in-figure text to publication-legible sizes and "
                "standardise across all figures, and (ii) a caption "
                "rewrite to make each caption self-contained.",
                False,
            ),
        ],
        [
            ("(a) Figure font-size standardisation. ", True),
            (
                "We audited every hardcoded ``fontsize=`` and "
                "``labelsize=`` value across the eight JRFM figure "
                "generators and found values as low as 8-11 pt, which "
                "rendered as sub-10 pt type when the figure was scaled to "
                "textwidth in an A4 layout. We applied a uniform size-"
                'bump rule (floor 12 pt, "+2" on moderate sizes) across '
                "all eight figures, producing a consistent typographic "
                "hierarchy (12 pt for smallest annotations, rising to "
                "16-18 pt for titles and display numbers). All six "
                "original figures (Figures 1-6) and the two revision-"
                "added figures (Figures 7-8) were regenerated from the "
                "bumped scripts. The one-shot bump script "
                "(``scripts/bump_font_sizes.py``) is committed in the "
                "code release so the change is reproducible.",
                False,
            ),
        ],
        [
            ("(b) Self-contained captions. ", True),
            (
                "Every caption now follows the rule: state (i) what is "
                "shown, (ii) the key numerical values a reader should "
                "notice, and (iii) what conclusion the reader should "
                "take from the figure. Five figure captions (Figures 1, "
                "3, 4, 5, 6) were rewritten in this pass; Figures 7 and "
                "8 and Tables 2-6 (added during other parts of the "
                "revision) were already written to this standard. Each "
                'rewritten caption ends with an explicit "Read this '
                'figure as:" clause giving the intended interpretation. '
                "For example, Figure 5 (GEX magnitude distribution) "
                'closes with "Read this figure as: the magnitude '
                "criterion alone -- before persistence or stability are "
                "even checked -- already separates the two eras, and "
                "the chosen $5B threshold is positioned in the trough "
                "between the two distributions rather than in the bulk "
                'of either."',
                False,
            ),
        ],
        [
            (
                "The resulting figures are materially easier to read at "
                "journal-print scale than the originals, with consistent "
                "font sizes across all eight figures and no in-figure "
                "text smaller than 12 pt.",
                False,
            ),
        ],
        [
            (
                "Change location: eight figure PNGs regenerated under docs/papers/paper2/figures/output/ and copied into docs/papers/jrfm/figures/; the bump script and all six modified figure generators committed; captions in Section 3 (Figure 1) and Section 4 (Figures 3, 4, 5, 6) rewritten in the .tex.",
                True,
            ),
        ],
    ],
}

ENGLISH_RESPONSE: list[list[tuple[str, bool]]] = [
    [
        ("We performed a full editing pass over the manuscript after all content changes were settled.", False),
    ],
    [
        (
            "We checked the manuscript for the usual English-editing "
            'offenders ("In order to", "It should be noted that", "It is '
            'worth noting", "Due to the fact that", "This is because", '
            '"Obviously", "Clearly"). None of these phrases appear in the '
            "manuscript; the original draft was already written in an "
            "active, direct register, so no changes were required on this "
            "axis.",
            False,
        ),
    ],
    [
        (
            "We identified the paragraphs with the most elaborate nested-"
            "clause sentences (the Section 1 philosophical opener and "
            "Section 5.5 Dispersed Knowledge) and rewrote them for "
            "directness. Section 1 was fully replaced in the rewrite for "
            "comment 1 above. Section 5.5 was tightened by breaking three "
            ">40-word sentences into two-sentence units while retaining the "
            "Hayek citation and the 30.8 pp empirical claim.",
            False,
        ),
    ],
    [
        (
            "We verified terminology consistency across sections: "
            '"regime" (not "state") for the detection target, '
            '"persistent / fragmented" (not "stable / unstable") for '
            'the binary outcome, "obfuscation" (not "anonymisation"), '
            '"dealer gamma positioning" where the detection task is the '
            "referent.",
            False,
        ),
    ],
    [
        (
            "Change location: targeted tightening in Section 6.5 Dispersed Knowledge; Section 1 opener and Section 6.3 Market Structure Evolution rewrites landed in the earlier comments 1 and 3 commits; technical-term consistency verified throughout.",
            True,
        ),
    ],
]


# ---------- Doc builder ----------

RED = RGBColor(0xC0, 0x00, 0x00)  # MDPI convention: responses in red
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_styled_paragraph(
    doc: Document, runs: list[tuple[str, bool]], *, response: bool, italic_first: bool = False
) -> None:
    p = doc.add_paragraph()
    for i, (text, bold) in enumerate(runs):
        run = p.add_run(text)
        run.bold = bold
        if response:
            run.font.color.rgb = RED
        if italic_first and i == 0:
            run.italic = True


def add_comment_block(doc: Document, n: str, comment: str, response_paragraphs: list[list[tuple[str, bool]]]) -> None:
    # "Comments N:"
    p = doc.add_paragraph()
    r = p.add_run(f"Comments {n}: ")
    r.bold = True
    r.font.color.rgb = BLACK
    p.add_run(comment)

    # "Response N:"
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Response {n}: ")
    r2.bold = True
    r2.font.color.rgb = RED
    # Mark revisions in red per template convention
    for run_set in response_paragraphs:
        add_styled_paragraph(doc, run_set, response=True)


def build() -> None:
    doc = Document()

    # Set default paragraph / font size
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Response to Reviewer 3 Comments", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "JRFM Submission jrfm-4256551 — Validating LLM Structural "
        "Reasoning: Detecting Persistent Market Regimes Through Temporal "
        "Obfuscation"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = doc.add_paragraph("Christopher Regan and Ying Xie, Kennesaw State University")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph("24 April 2026")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Summary
    doc.add_heading("1. Summary", level=1)
    doc.add_paragraph(
        "Thank you very much for taking the time to review this "
        "manuscript and for the substantive, constructive feedback. The "
        "reviewer's comments identified meaningful improvements in "
        "introduction focus, contribution positioning, benchmark "
        "comparison, methodological transparency, statistical rigour, "
        "practitioner connection, limitations scope, and figure clarity. "
        "We have addressed every point-by-point comment in the revised "
        "manuscript; detailed responses and the corresponding revisions "
        "(marked in red) are provided below. The revised manuscript is "
        "31 A4 pages, up from 18 in the originally submitted version."
    )
    # The reviewer's own summary paragraph
    doc.add_paragraph(f"Reviewer's summary: “{REVIEWER_SUMMARY_COMMENT}”")

    # 2. Questions for General Evaluation
    doc.add_heading("2. Questions for General Evaluation", level=1)
    rows = [
        (
            "Does the introduction provide sufficient background and include all relevant references?",
            "Must be improved",
            'Addressed in Comment 1 below: Section 1 rewritten, new "Research gap" paragraph, new 2022-2025 references (Dim, Eraker & Vilkov 2023).',
        ),
        (
            "Is the research design appropriate?",
            "Can be improved",
            "Addressed in Comment 3: new Markov-switching benchmark (Section 3.8 + Section 5.6) demonstrates the framework is not reducible to a volatility-regime classifier.",
        ),
        (
            "Are the methods adequately described?",
            "Can be improved",
            "Addressed in Comment 4: new Appendix A reproduces the full LLM prompt verbatim; new Section 5.5 reports a 45-configuration threshold-sensitivity sweep; reproducibility posture documented.",
        ),
        (
            "Are the results clearly presented?",
            "Can be improved",
            "Addressed in Comment 5: every detection rate in Section 4 now carries a 95% bootstrap or Wilson CI; full chi-square / Fisher statistics reported; robustness confirmed.",
        ),
        (
            "Are the conclusions supported by the results?",
            "Can be improved",
            "Addressed in Comments 2 and 5(d): positioning statement added to Section 1.3 and echoed in Section 7 opening; strong-claim language moderated where CIs or sensitivity warranted.",
        ),
        (
            "Are all figures and tables clear and well-presented?",
            "Can be improved",
            'Addressed in Comment 8: captions on Figures 1, 3, 4, 5, 6 rewritten to be self-contained with explicit "Read this figure as:" clauses.',
        ),
        (
            "Quality of English Language",
            "Can be improved",
            "Addressed in Section 4 of this response: targeted tightening in Section 5.5; wordy-transition check passed; terminology consistency verified.",
        ),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(("Question", "Reviewer's evaluation", "Response and revisions")):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(text)
        run.bold = True
    for i, (q, ev, resp) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = q
        cells[1].text = ev
        cells[2].text = resp
        for c in cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 3. Point-by-point response
    doc.add_heading("3. Point-by-point response to Comments and Suggestions for Authors", level=1)
    for n, comment in COMMENTS:
        add_comment_block(doc, n, comment, R[n])

    # 4. English Language
    doc.add_heading("4. Response to Comments on the Quality of English Language", level=1)
    # Reviewer comment
    p = doc.add_paragraph()
    r = p.add_run("Point 1: ")
    r.bold = True
    p.add_run(ENGLISH_COMMENT)
    # Response
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Response 1: ")
    r2.bold = True
    r2.font.color.rgb = RED
    for run_set in ENGLISH_RESPONSE:
        add_styled_paragraph(doc, run_set, response=True)

    # 5. Additional clarifications
    doc.add_heading("5. Additional clarifications", level=1)
    doc.add_paragraph(
        "We have also raised with the handling editor (separately, via "
        "the portal comments-to-editor field) that Reviewer 1's report "
        "appears to apply to a different manuscript — it asks about "
        "conformable derivatives in the Heston framework, Heston-He-Zhu "
        "comparisons, jump-diffusion and fractional models, and "
        "computational challenges in an option-pricing algorithm, none "
        "of which are topics our manuscript addresses. We are prepared "
        "to respond substantively once the correct review is available, "
        "or to a replacement reviewer if that is more expedient. This "
        "clarification is orthogonal to Reviewer 3's comments and does "
        "not affect the revisions above."
    )
    doc.add_paragraph(
        "We have also incorporated Reviewer 2's recommendation for "
        "acceptance (uploaded separately as the Reviewer 2 response)."
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
